#!/usr/bin/python

import sqlite3
from shutil import copyfile
import pandas as pd
import os
import json
from json2xml import json2xml
from json2xml.utils import readfromjson
import asyncio
from PIL import Image
from pyppeteer import launch
from datetime import datetime
import glob
import shutil
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_DIRECTION, WD_TABLE_ALIGNMENT
import pyexcel as pe
import csv
from config import *


def init():
    if not os.path.exists('TEMP'):
        os.makedirs('TEMP')

    if not os.path.exists(RELEASE_DIR):
        os.makedirs(RELEASE_DIR)

    if not os.path.exists(RELEASE_DIR + '/CSV'):
        os.makedirs(RELEASE_DIR + '/CSV')

    if not os.path.exists(RELEASE_DIR + '/docs'):
        os.makedirs(RELEASE_DIR + '/docs')

    if not os.path.exists(RELEASE_DIR + '/Excel'):
        os.makedirs(RELEASE_DIR + '/Excel')

    if not os.path.exists(RELEASE_DIR + '/Flash Card'):
        os.makedirs(RELEASE_DIR + '/Flash Card')

    if not os.path.exists(RELEASE_DIR + '/JSON'):
        os.makedirs(RELEASE_DIR + '/JSON')

    if not os.path.exists(RELEASE_DIR + '/PDF'):
        os.makedirs(RELEASE_DIR + '/PDF')

    if not os.path.exists(RELEASE_DIR + '/SQLite'):
        os.makedirs(RELEASE_DIR + '/SQLite')

    if not os.path.exists(RELEASE_DIR + '/Word'):
        os.makedirs(RELEASE_DIR + '/Word')

    if not os.path.exists(RELEASE_DIR + '/XML'):
        os.makedirs(RELEASE_DIR + '/XML')

    copyfile("./templates/favicon.png", RELEASE_DIR + "/docs/favicon.png")
    copyfile("./templates/index.html", RELEASE_DIR + "/docs/index.html")
    copyfile("./templates/LICENSE", RELEASE_DIR + "/LICENSE")


def normalize(path: str = RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite"):
    try:
        conn = sqlite3.connect(DB_DIR)
        cursor = conn.cursor()
        print("Connected to SQLite successfully .")
        print("Normalization started...")
        sqlite_update_query = "UPDATE tbl_words SET glk_word = trim(glk_word), glk_example = trim(glk_example), " \
                              "fa_word = trim(fa_word), fa_example = trim(fa_example), en_word = trim(en_word), " \
                              "en_example = trim(en_example) "
        cursor.execute(sqlite_update_query)
        conn.commit()
        print("Records striped successfully.", cursor.rowcount)
        cursor.close()

        print("Releasing SQLite...")
        copyfile(DB_DIR, RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite")
        print("SQLite Released Successfully.")

    except sqlite3.Error as error:
        print("Failed to stripe rows.", error)
    finally:
        if (conn):
            conn.close()
            print("The SQLite connection is closed.")

    print("Sorting database started...")
    copyfile("./templates/database.db", path)
    try:
        conn = sqlite3.connect(DB_DIR)
        cursor = conn.cursor()

        sqlite_select_query = "SELECT * from tbl_words order by glk_word"
        cursor.execute(sqlite_select_query)
        records = cursor.fetchall()

        i = 0
        conn2 = sqlite3.connect(path)
        for row in records:
            i += 1
            cursor2 = conn2.cursor()
            r1 = row[1]  # glk_word
            r2 = row[2]  # glk_example
            r3 = row[3]  # fa_word
            r4 = row[4]  # fa_example
            r5 = row[5]  # en_word
            r6 = row[6]  # en_example
            r1 = r1.replace("'", "ٰ")
            r2 = r2.replace("'", "ٰ")
            r3 = r3.replace("'", "ٰ")
            r3 = r3.replace("ي", "ی")
            r4 = r4.replace("'", "ٰ")
            r4 = r4.replace("ي", "ی")
            r5 = r5.replace("ٰ", "'")
            r6 = r6.replace("ٰ", "'")
            sqlite_insert_query = f"""INSERT INTO tbl_words
                                      (id, glk_word, glk_example, fa_word, fa_example, en_word, en_example) 
                                       VALUES 
                                      ({i}, "{r1}", "{r2}", "{r3}", "{r4}", "{r5}", "{r6}")"""

            cursor2.execute(sqlite_insert_query)
            conn2.commit()
        cursor.close()

    except sqlite3.Error as error:
        print("Failed to read data from sqlite table", error)
    finally:
        if (conn):
            conn.close()
            print("Database sorted Successfully.")


def make_csv():
    print("Making CSV started...")

    conn = sqlite3.connect(RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite", isolation_level=None,
                           detect_types=sqlite3.PARSE_COLNAMES)

    tmp_output_dir = RELEASE_DIR + "/CSV/Top 1000 Words in Gilaki (Full).csv"
    db_df = pd.read_sql_query("SELECT * FROM tbl_words", conn)
    db_df.to_csv(tmp_output_dir, index=False, quoting=csv.QUOTE_NONNUMERIC)

    tmp_output_dir = RELEASE_DIR + "/CSV/Top 1000 Words in Gilaki (Farsi Only).csv"
    db_df = pd.read_sql_query("SELECT id, glk_word, glk_example, fa_word, fa_example FROM tbl_words", conn)
    db_df.to_csv(tmp_output_dir, index=False, quoting=csv.QUOTE_NONNUMERIC)

    tmp_output_dir = RELEASE_DIR + "/CSV/Top 1000 Words in Gilaki (English Only).csv"
    db_df = pd.read_sql_query("SELECT id, glk_word, glk_example, en_word, en_example FROM tbl_words", conn)
    db_df.to_csv(tmp_output_dir, index=False, quoting=csv.QUOTE_NONNUMERIC)

    print("CSV made Successfully.")


def make_json():
    print("Making JSON started...")

    conn = sqlite3.connect(RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite")
    conn.row_factory = sqlite3.Row
    db = conn.cursor()
    rows = db.execute("SELECT * from tbl_words").fetchall()
    conn.commit()
    conn.close()
    js = json.dumps([dict(ix) for ix in rows], ensure_ascii=False).encode('utf8')
    tmp_min_output_dir = RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki.min.json"
    f = open(tmp_min_output_dir, "w")
    f.write(js.decode())
    f.close()
    tmp_output_dir = RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki.json"
    js = json.dumps([dict(ix) for ix in rows], indent=4, ensure_ascii=False).encode('utf8')
    f = open(tmp_output_dir, "w")
    f.write(js.decode())
    f.close()

    conn = sqlite3.connect(RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite")
    conn.row_factory = sqlite3.Row
    db = conn.cursor()
    rows = db.execute("SELECT id, glk_word, glk_example, fa_word, fa_example from tbl_words").fetchall()
    conn.commit()
    conn.close()
    js = json.dumps([dict(ix) for ix in rows], ensure_ascii=False).encode('utf8')
    tmp_min_output_dir = RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki (Farsi Only).min.json"
    f = open(tmp_min_output_dir, "w")
    f.write(js.decode())
    f.close()
    tmp_output_dir = RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki (Farsi Only).json"
    js = json.dumps([dict(ix) for ix in rows], indent=4, ensure_ascii=False).encode('utf8')
    f = open(tmp_output_dir, "w")
    f.write(js.decode())
    f.close()

    conn = sqlite3.connect(RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite")
    conn.row_factory = sqlite3.Row
    db = conn.cursor()
    rows = db.execute("SELECT id, glk_word, glk_example, en_word, en_example from tbl_words").fetchall()
    conn.commit()
    conn.close()
    js = json.dumps([dict(ix) for ix in rows], ensure_ascii=False).encode('utf8')
    tmp_min_output_dir = RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki (English Only).min.json"
    f = open(tmp_min_output_dir, "w")
    f.write(js.decode())
    f.close()
    tmp_output_dir = RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki (English Only).json"
    js = json.dumps([dict(ix) for ix in rows], indent=4, ensure_ascii=False).encode('utf8')
    f = open(tmp_output_dir, "w")
    f.write(js.decode())
    f.close()

    print("JSON made Successfully.")


def update_docs():
    print("Updating docs...")
    copyfile(RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki.min.json", RELEASE_DIR + "/docs/Top 1000 Words in "
                                                                                    "Gilaki.min.json")
    print("docs updated Successfully.")


def make_xml():
    print("Making XML started...")

    tmp_output_dir = RELEASE_DIR + "/XML/Top 1000 Words in Gilaki.xml"
    data = readfromjson(RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki.min.json")
    data = json2xml.Json2xml(data, wrapper="wordlist", pretty=True, attr_type=False).to_xml()
    data = data.replace("<item>", "<word>")
    data = data.replace("</item>", "</word>")
    data = data.replace('<?xml version="1.0" ?>', '<?xml version="1.0" encoding="UTF-8"?>')
    f = open(tmp_output_dir, "w")
    f.write(data)
    f.close()

    tmp_output_dir = RELEASE_DIR + "/XML/Top 1000 Words in Gilaki (Farsi Only).xml"
    data = readfromjson(RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki (Farsi Only).min.json")
    data = json2xml.Json2xml(data, wrapper="wordlist", pretty=True, attr_type=False).to_xml()
    data = data.replace("<item>", "<word>")
    data = data.replace("</item>", "</word>")
    data = data.replace('<?xml version="1.0" ?>', '<?xml version="1.0" encoding="UTF-8"?>')
    f = open(tmp_output_dir, "w")
    f.write(data)
    f.close()

    tmp_output_dir = RELEASE_DIR + "/XML/Top 1000 Words in Gilaki (English Only).xml"
    data = readfromjson(RELEASE_DIR + "/JSON/Top 1000 Words in Gilaki (English Only).min.json")
    data = json2xml.Json2xml(data, wrapper="wordlist", pretty=True, attr_type=False).to_xml()
    data = data.replace("<item>", "<word>")
    data = data.replace("</item>", "</word>")
    data = data.replace('<?xml version="1.0" ?>', '<?xml version="1.0" encoding="UTF-8"?>')
    f = open(tmp_output_dir, "w")
    f.write(data)
    f.close()

    print("XML made Successfully.")


def readtemplate(path: str) -> str:
    f = open(f"{path}", "r")
    return f.read()


def make_flash_html(path: str = RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite"):
    print("Making HTML flashcards started...")

    word = readtemplate("./templates/word.html")
    meaning = readtemplate("./templates/meaning.html")
    full = readtemplate("./templates/full_fa.html")
    full_en = readtemplate("./templates/full_en.html")

    try:
        conn = sqlite3.connect(path)
        cursor = conn.cursor()

        sqlite_select_query = "SELECT * from tbl_words"
        cursor.execute(sqlite_select_query)
        records = cursor.fetchall()

        for row in records:
            flash = word
            flash = flash.replace("{glk_word}", row[1])
            createflash(f"{row[0]}-word.html", flash)

        for row in records:
            flash = meaning
            flash = flash.replace("{glk_word}", row[1])
            flash = flash.replace("{glk_example}", row[2])
            createflash(f"{row[0]}-meaning.html", flash)

        for row in records:
            txt = ""
            flash = full
            flash = flash.replace("{glk_word}", row[1])
            txt += row[1] + "\n"
            flash = flash.replace("{fa_word}", row[3])
            txt += row[3] + "\n"
            flash = flash.replace("{glk_example}", row[2])
            txt += row[2] + "\n"
            flash = flash.replace("{fa_example}", row[4])
            txt += row[4] + "\n"
            createflash(f"{row[0]}-full-fa.html", flash)
            txt += row[5] + "\n"
            txt += row[6] + "\n"
            createinfo(f"{row[0]}.txt", txt)

        for row in records:
            txt = ""
            flash = full_en
            flash = flash.replace("{glk_word}", row[1])
            txt += row[1] + "\n"
            flash = flash.replace("{en_word}", row[5])
            txt += row[3] + "\n"
            flash = flash.replace("{glk_example}", row[2])
            txt += row[2] + "\n"
            flash = flash.replace("{en_example}", row[6])
            txt += row[4] + "\n"
            createflash(f"{row[0]}-full-en.html", flash)

        cursor.close()

    except sqlite3.Error as error:
        print("Failed to read data from sqlite table", error)
    finally:
        if (conn):
            conn.close()
            print("HTML flashcards made Successfully.")


def crop_center(pil_img, crop_width, crop_height):
    img_width, img_height = pil_img.size
    return pil_img.crop(((img_width - crop_width) // 2,
                         (img_height - crop_height) // 2,
                         (img_width + crop_width) // 2,
                         (img_height + crop_height) // 2))


def makejpg(typ):
    for i in range(START_RANGE, END_RANGE, 1):
        async def main():
            browser = await launch(headless=True, options={'args': ['--no-sandbox']})
            page = await browser.newPage()
            await page.setViewport({'width': 600, 'height': 600})
            p = os.path.abspath(f"TEMP/{i}-{typ}.html")
            await page.goto(f"file://{p}")
            await page.screenshot({'path': f"TEMP/{i}-{typ}.jpg"})
            await browser.close()

        asyncio.get_event_loop().run_until_complete(main())

        thumb_width = 600

        im = Image.open(f"TEMP/{i}-{typ}.jpg")
        im_thumb = crop_center(im, thumb_width, thumb_width)
        im_thumb.save(f"TEMP/{i}-{typ}.jpg", quality=100)


def createflash(name, html):
    f = open(f"TEMP/{name}", "w")
    f.write(html)
    f.close()


def createinfo(name, text):
    f = open(f"TEMP/{name}", "w")
    f.write(text)
    f.close()


def make_flash_jpg():
    print("Making JPG flashcards started...")

    makejpg("word")
    makejpg("meaning")
    makejpg("full-fa")
    makejpg("full-en")

    print("JPG flashcards made Successfully.")

    print("Releasing JPG flashcards...")
    files = glob.iglob(os.path.join("./TEMP/", "*.jpg"))
    for file in files:
        if os.path.isfile(file):
            shutil.copy2(file, RELEASE_DIR + "/Flash Card/")
    print("JPG flashcards Released Successfully.")


def make_gif(lang_code: str):
    print(f"Making {lang_code} GIF flashcards started...")

    def makegif(num):
        os.system(
            f'convert -delay 200 -loop 0 ./TEMP/{num}-word.jpg ./TEMP/{num}-meaning.jpg ./TEMP/{num}-full-{lang_code}.jpg ./TEMP/{num}-animation-{lang_code}.gif')

    for i in range(START_RANGE, END_RANGE, 1):
        makegif(i)

    print(f"{lang_code} GIF flashcards made Successfully.")
    print(f"Releasing {lang_code} GIF flashcards...")
    files = glob.iglob(os.path.join("./TEMP/", "*.gif"))
    for file in files:
        if os.path.isfile(file):
            shutil.copy2(file, RELEASE_DIR + "/Flash Card/")
            os.remove(file)
    print(f"{lang_code} GIF flashcards Released Successfully.")


def make_docx(lang_code: str = "fa", path: str = RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite"):
    language_name = "Farsi"

    if lang_code == "en":
        language_name = "English"

    print(f"Making {language_name} DOCX started...")

    document = Document("./templates/database.docx")
    table = document.tables[0]
    table.directions = WD_TABLE_DIRECTION.RTL
    table.allow_autofit = True

    row = table.rows[0].cells

    row[0].text = ''
    if lang_code == "fa":
        row[0].paragraphs[0].add_run(f'مثال فارسی').bold = True
    if lang_code == "en":
        row[0].paragraphs[0].add_run(f'English Example').bold = True
    row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    row[1].text = ''
    if lang_code == "fa":
        row[1].paragraphs[0].add_run(f'کلمه فارسی').bold = True
    if lang_code == "en":
        row[1].paragraphs[0].add_run(f'English Word').bold = True
    row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    row[2].text = ''
    if lang_code == "fa":
        row[2].paragraphs[0].add_run('مثال گیلکی').bold = True
    if lang_code == "en":
        row[2].paragraphs[0].add_run('Gilaki Example').bold = True
    row[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    row[3].text = ''
    if lang_code == "fa":
        row[3].paragraphs[0].add_run('کلمه گیلکی').bold = True
    if lang_code == "en":
        row[3].paragraphs[0].add_run('Gilaki Word').bold = True
    row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    try:
        conn = sqlite3.connect(path)
        cursor = conn.cursor()

        sqlite_select_query = f"SELECT id, glk_word, glk_example, {lang_code}_word, {lang_code}_example from tbl_words"
        cursor.execute(sqlite_select_query)
        records = cursor.fetchall()

        for r in records:
            row = table.add_row().cells

            row[0].text = ''
            row[0].paragraphs[0].add_run(r[4])
            row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row[0].alignment = WD_TABLE_ALIGNMENT.RIGHT

            row[1].text = ''
            row[1].paragraphs[0].add_run(r[3])
            row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row[1].alignment = WD_TABLE_ALIGNMENT.RIGHT

            row[2].text = ''
            row[2].paragraphs[0].add_run(r[2])
            row[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row[2].alignment = WD_TABLE_ALIGNMENT.RIGHT

            row[3].text = ''
            row[3].paragraphs[0].add_run(r[1])
            row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row[3].alignment = WD_TABLE_ALIGNMENT.RIGHT

        cursor.close()

    except sqlite3.Error as error:
        print("Failed to read data from sqlite table", error)
    finally:
        if (conn):
            conn.close()
            print(f"{language_name} DOCX made Successfully.")

    document.save(RELEASE_DIR + f"/Word/Top 1000 Words in Gilaki ({language_name}).docx")


def make_pdf(lang_code: str = "fa"):
    language_name = "Farsi"

    if lang_code == "en":
        language_name = "English"

    print(f"Making {language_name} PDF started...")

    tmp_dir = RELEASE_DIR + f"/Word/Top 1000 Words in Gilaki ({language_name}).docx"
    os.system(f'libreoffice --headless --convert-to pdf --outdir "{RELEASE_DIR}/PDF/" "{tmp_dir}"')

    print(f"{language_name} PDF made Successfully.")


def make_xlsx(lang_code: str = "fa", path: str = RELEASE_DIR + "/SQLite/Top 1000 Words in Gilaki.sqlite"):
    language_name = "Farsi"

    if lang_code == "en":
        language_name = "English"

    print(f"Making {language_name} Excel started...")

    sheet = pe.get_sheet(file_name=f"./templates/database{lang_code}.xlsx")

    try:
        conn = sqlite3.connect(path)
        cursor = conn.cursor()

        sqlite_select_query = f"SELECT id, glk_word, glk_example, {lang_code}_word, {lang_code}_example from tbl_words"
        cursor.execute(sqlite_select_query)
        records = cursor.fetchall()

        for r in records:
            sheet.row += [r[1], r[2], r[3], r[4]]

        cursor.close()

        sheet.rightToLeft = True
        sheet.save_as(RELEASE_DIR + f"/Excel/Top 1000 Words in Gilaki ({language_name}).xlsx")

    except sqlite3.Error as error:
        print("Failed to read data from sqlite table", error)
    finally:
        if (conn):
            conn.close()
            print("XLSX made Successfully.")


def change_readme():
    print("Updating README started...")
    tmp_dir = "./templates/README.txt"
    readme = readtemplate(tmp_dir)
    d = datetime.now()
    current_time = d.strftime('%d, %b %Y')
    readme = readme.replace("{date}", f"{current_time}")
    readme = readme.replace("{version}", f"{VERSION}")
    f = open(RELEASE_DIR + "/README.md", "w")
    f.write(readme)
    f.close()
    print("README updated Successfully.")


def push_release():
    os.system(f'cd {RELEASE_DIR} && git add . && git commit -m "{COMMIT_MESSAGE}" && git push')


if __name__ == '__main__':
    init()
    normalize()
    make_csv()
    make_json()
    update_docs()
    make_xml()
    make_docx("fa")
    make_docx("en")
    make_pdf("fa")
    make_pdf("en")
    make_xlsx("fa")
    make_xlsx("en")
    make_flash_html()
    make_flash_jpg()
    make_gif("fa")
    make_gif("en")
    change_readme()
    push_release()
